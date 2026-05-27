"""
DOCX adapter.

The challenge: detection runs on a flat plaintext string, but a .docx stores
text in a tree of paragraphs -> runs (a "run" is a span of uniform formatting).
To edit without destroying formatting we must:

  1. Walk every run in document order, concatenating run.text into one string
     while recording, for each run, its [start, end) offset in that string.
     (We join paragraphs with "\n" so offsets line up with what the user sees.)

  2. After detection produces edits as (start, end, replacement) on the flat
     string, map each edit back onto the run(s) it touches and rewrite those
     runs' .text in place. Formatting on each run is preserved because we only
     touch the text content, not the run's style.

Edits that span multiple runs are handled by writing the full replacement into
the first overlapped run and clearing the remainder of the overlap in the others.
"""
from __future__ import annotations

import io
from typing import List, Tuple

from docx import Document


def _iter_runs(doc):
    """Yield every run in body paragraphs and table cells, in reading order."""
    def runs_in_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                yield r
            yield None  # paragraph break marker

    yield from runs_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from runs_in_paragraphs(cell.paragraphs)


def _build_run_map(doc):
    """Return (flat_text, [(run, start, end), ...]) — None entries are breaks."""
    text_parts = []
    run_spans = []  # (run_or_None, start, end)
    cursor = 0
    for item in _iter_runs(doc):
        if item is None:  # paragraph break
            text_parts.append("\n")
            run_spans.append((None, cursor, cursor + 1))
            cursor += 1
        else:
            t = item.text or ""
            text_parts.append(t)
            run_spans.append((item, cursor, cursor + len(t)))
            cursor += len(t)
    return "".join(text_parts), run_spans


# We carry the parsed Document between extract and reassemble via a small state
# object so we don't parse the bytes twice.
class DocxState:
    def __init__(self, raw_bytes: bytes):
        self.doc = Document(io.BytesIO(raw_bytes))
        self.text, self.run_spans = _build_run_map(self.doc)


def extract(raw_bytes: bytes) -> Tuple[str, "DocxState"]:
    state = DocxState(raw_bytes)
    return state.text, state


def reassemble(state: "DocxState", edits: List[Tuple[int, int, str]]) -> bytes:
    """Apply edits to the runs and serialize the docx back to bytes."""
    edits = sorted(edits, key=lambda e: e[0])

    # For each run, collect the slice edits that fall within it.
    # Build new text per run by walking its own [rs, re) window.
    for run, rs, re_ in state.run_spans:
        if run is None:
            continue
        run_text = run.text or ""
        # Find edits overlapping this run.
        new_text = []
        local_cursor = rs
        for start, end, repl in edits:
            if end <= rs or start >= re_:
                continue  # no overlap with this run
            # clamp edit to run boundaries
            cs = max(start, rs)
            ce = min(end, re_)
            # text before the edit (in this run)
            new_text.append(run_text[local_cursor - rs:cs - rs])
            # write the replacement only in the run where the edit STARTS,
            # so a multi-run edit isn't duplicated.
            if start >= rs:
                new_text.append(repl)
            local_cursor = ce
        new_text.append(run_text[local_cursor - rs:])
        candidate = "".join(new_text)
        if candidate != run_text:
            run.text = candidate

    buf = io.BytesIO()
    state.doc.save(buf)
    return buf.getvalue()
